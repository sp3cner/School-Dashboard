"""
Best-effort extraction of "important date" rows from an uploaded syllabus
(PDF or DOCX). Syllabi are not structured data, so this uses heuristics:

  1. Pull raw text out of the file.
  2. Scan line-by-line (and small windows of lines) for a date.
  3. If a line has a date, keep the surrounding text as the "item" label,
     and tag it with a rough category (exam / paper / project / reading / other)
     based on keywords.

This will not be perfect for every syllabus format. The dashboard lets the
user review and delete/edit rows before they're merged into the timeline.
"""

from __future__ import annotations
import re
import io
from datetime import datetime
from typing import Optional
from dateutil import parser as dateparser

import pdfplumber
import docx


DATE_PATTERNS = [
    # Month name + day (+ optional year): "September 5", "Sept. 5, 2026"
    r"\b(?:Jan(?:uary)?|Feb(?:ruary)?|Mar(?:ch)?|Apr(?:il)?|May|Jun(?:e)?|Jul(?:y)?|"
    r"Aug(?:ust)?|Sep(?:t(?:ember)?)?|Oct(?:ober)?|Nov(?:ember)?|Dec(?:ember)?)\.?\s+\d{1,2}(?:st|nd|rd|th)?(?:,?\s*\d{4})?\b",
    # numeric dates: 9/5, 9/5/26, 9-5-2026
    r"\b\d{1,2}[/-]\d{1,2}(?:[/-]\d{2,4})?\b",
]
DATE_RE = re.compile("|".join(DATE_PATTERNS), re.IGNORECASE)

KEYWORD_CATEGORY = {
    "exam": ["exam", "midterm", "final"],
    "quiz": ["quiz"],
    "paper": ["paper", "essay"],
    "project": ["project"],
    "presentation": ["presentation"],
    "reading": ["read", "chapter"],
    "assignment": ["assignment", "homework", "hw", "due"],
    "class_event": ["no class", "holiday", "break"],
}


def _categorize(text: str) -> str:
    lowered = text.lower()
    for category, keywords in KEYWORD_CATEGORY.items():
        if any(k in lowered for k in keywords):
            return category
    return "other"


def extract_text(file_bytes: bytes, filename: str) -> str:
    ext = filename.lower().rsplit(".", 1)[-1]
    if ext == "pdf":
        text_parts = []
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text() or ""
                text_parts.append(page_text)
        return "\n".join(text_parts)
    elif ext in ("docx",):
        document = docx.Document(io.BytesIO(file_bytes))
        lines = [p.text for p in document.paragraphs]
        # also pull table cells - many syllabi put the schedule in a table
        for table in document.tables:
            for row in table.rows:
                lines.append(" | ".join(cell.text for cell in row.cells))
        return "\n".join(lines)
    else:
        raise ValueError(f"Unsupported file type: .{ext}. Use PDF or DOCX.")


def extract_dated_items(text: str, source_filename: str, default_year: Optional[int] = None) -> list[dict]:
    """Scan text line by line for anything that looks like a date, and
    return a list of {date, raw_line, category, source_file} rows."""
    rows = []
    now_year = default_year or datetime.now().year

    for raw_line in text.splitlines():
        line = raw_line.strip()
        if not line or len(line) < 4:
            continue
        matches = DATE_RE.findall(line)
        if not matches:
            continue
        for match in matches:
            try:
                parsed = dateparser.parse(match, default=datetime(now_year, 1, 1), fuzzy=True)
            except (ValueError, OverflowError):
                continue
            rows.append({
                "source": "Syllabus",
                "source_file": source_filename,
                "date": parsed.date().isoformat(),
                "item": line[:200],
                "category": _categorize(line),
            })
    return rows


def parse_syllabus_file(file_bytes: bytes, filename: str, default_year: Optional[int] = None) -> list[dict]:
    text = extract_text(file_bytes, filename)
    return extract_dated_items(text, filename, default_year=default_year)
